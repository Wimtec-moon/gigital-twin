# Aktivieren der Widgets-Erweiterung
import warnings
warnings.filterwarnings('ignore')

!jupyter nbextension enable --py --sys-prefix widgetsnbextension

!pip install ipywidgets==7.7.2

!pip install notebook
!pip install ipywidgets
!pip install widgetsnbextension

!jupyter nbextension enable --py widgetsnbextension --sys-prefix
!jupyter nbextension enable --py ipywidgets --sys-prefix

!pip install --upgrade ipywidgets

!pip install python-docx

import plotly.express as px
import ipywidgets as widgets
from IPython.display import display, clear_output
import os
import pandas as pd
from docx import Document
from docx.shared import Inches
import io
import matplotlib.pyplot as plt
from IPython.display import Image, display, HTML

# CI-Farben
primary_color = "#001c2f"  # Night Sky
secondary_color = "#f0f0f0"  # Shade One
accent_color = "#9dfdb4"  # Mint Energy

# Funktion zum Laden des Mutterdatensatzes aus dem ursprünglichen Speicherort
def load_mother_data():
    file_path = 'hyperlogmotherdata.csv'  # Passen Sie den Pfad zu Ihrer CSV-Datei an
    if os.path.exists(file_path):
        df = pd.read_csv(
            file_path,
            sep=';',  # Passen Sie das Trennzeichen an
            encoding='ISO-8859-1',  # Kodierung auf 'ISO-8859-1' setzen
            header=0  # Kopfzeile einlesen
        )
        # Spaltennamen säubern (Leerzeichen entfernen)
        df.columns = df.columns.str.strip()
        # Spaltennamen anzeigen
        #print("Spalten im DataFrame:", df.columns.tolist())

        # Umbenennen der Spalten, falls erforderlich
        if 'Start Time' not in df.columns and 'Startzeit' in df.columns:
            df.rename(columns={'Startzeit': 'Start Time'}, inplace=True)
        if 'End Time' not in df.columns and 'Endzeit' in df.columns:
            df.rename(columns={'Endzeit': 'End Time'}, inplace=True)
        if 'Charger Serial Number' not in df.columns and 'Seriennummer' in df.columns:
            df.rename(columns={'Seriennummer': 'Charger Serial Number'}, inplace=True)
        if 'Country' not in df.columns and 'Land' in df.columns:
            df.rename(columns={'Land': 'Country'}, inplace=True)
        if 'HYC Errorcode' not in df.columns and 'HYC Fehlercode' in df.columns:
            df.rename(columns={'HYC Fehlercode': 'HYC Errorcode'}, inplace=True)
        if 'SoC Start (%)' not in df.columns and 'SoC Start' in df.columns:
            df.rename(columns={'SoC Start': 'SoC Start (%)'}, inplace=True)
        if 'SoC Stop (%)' not in df.columns and 'SoC Stop' in df.columns:
            df.rename(columns={'SoC Stop': 'SoC Stop (%)'}, inplace=True)
        if 'Chargingsession ID' not in df.columns and 'Ladevorgangs-ID' in df.columns:
            df.rename(columns={'Ladevorgangs-ID': 'Chargingsession ID'}, inplace=True)

        # Konvertieren der Datumsfelder
        if 'Start Time' in df.columns:
            df['Start Time'] = pd.to_datetime(df['Start Time'], errors='coerce', dayfirst=True)
        if 'End Time' in df.columns:
            df['End Time'] = pd.to_datetime(df['End Time'], errors='coerce', dayfirst=True)

        return df
    else:
        print(f"Die Datei unter {file_path} wurde nicht gefunden.")
        return pd.DataFrame()

    
# **HINZUGEFÜGT:** Globale Variablen für Abteilungen und Subkarten
abteilungen = ['Vertrieb', 'Produktmanagement', 'Marketing', 'After Sales', 'Datenpflege', 'Reporting']
subkarten = {
    'Vertrieb': ['Chargerausnutzung', 'Chargeranalyse', 'Standortanalyse', 'Fahrzeuganalyse', 'Amortisationsrechnung'],
    'Produktmanagement': ['Produktausnutzung', 'Nutzeranalyse', 'Ländervergleich', 'Performance'],
    'Marketing': ['Kundenanalyse', 'Länderanalyse'],
    'After Sales': ['Fehlerquote', 'Fehlerarten', 'Charger Fehleranalyse'],
    'Datenpflege': ['Datenpflege'],
    'Reporting': ['Reporting']
}
# Mutterdatensatz laden
df_global = load_mother_data()

# Überprüfen, ob 'Start Time' im DataFrame vorhanden ist
if 'Start Time' not in df_global.columns:
    raise ValueError("Die Spalte 'Start Time' ist nicht im DataFrame vorhanden. Überprüfen Sie das Trennzeichen und die Kodierung Ihrer CSV-Datei.")

# Ausgabe der Spaltennamen zur Überprüfung
#print("Spaltennamen im DataFrame:", df_global.columns.tolist())

# Globale Funktion zur Aktualisierung des Dashboards (Seite neu laden)
def refresh_dashboard():
    clear_output(wait=True)
    start_menu()

# Funktion zur Filterung der Daten
def filter_data(df, seriennummer=None, start_date=None, end_date=None):
    filtered_df = df.copy()
    if seriennummer:
        filtered_df = filtered_df[filtered_df['Charger Serial Number'].astype(str) == seriennummer]
    if start_date:
        filtered_df = filtered_df[filtered_df['Start Time'] >= pd.to_datetime(start_date)]
    if end_date:
        filtered_df = filtered_df[filtered_df['Start Time'] <= pd.to_datetime(end_date)]
    return filtered_df
# Fahrzeuganalyse
import plotly.express as px
import ipywidgets as widgets
from IPython.display import display, clear_output
import pandas as pd
import matplotlib.pyplot as plt
import numpy as np

# CI-Farben
primary_color = "#001c2f"  # Night Sky
secondary_color = "#f0f0f0"  # Shade One
accent_color = "#9dfdb4"  # Mint Energy

# Beispiel-Daten für Dokumente
documents_data = {
    'Name': ['Datenblatt HYC50', 'Handbuch C7', 'Verkaufsunterlage MOON Power'],
    'Typ': ['Datenblatt', 'Handbuch', 'Verkaufsunterlage'],
    'Pfad': ['path/to/hyc50_datenblatt.pdf', 'path/to/c7_handbuch.pdf', 'path/to/moon_power_verkaufsunterlage.pdf']
}
documents_df = pd.DataFrame(documents_data)

# Fahrzeuganalyse-Dashboard
def fahrzeuganalyse_dashboard(df):
    container = widgets.Output()

    def plot_vehicle_analysis(filtered_df):
        with container:
            container.clear_output()

            # Überprüfen, ob die Spalte 'Car' vorhanden ist
            if 'Car' not in filtered_df.columns:
                display("Die Spalte 'Car' ist nicht im DataFrame vorhanden.")
                return

            # Schritt 1: Fahrzeugtypen zählen und prozentuale Verteilung berechnen
            vehicle_counts = filtered_df['Car'].value_counts()
            vehicle_percentages = vehicle_counts / vehicle_counts.sum() * 100

            # Fahrzeugtypen unter 2% als 'Sonstiges' zusammenfassen
            other_threshold = 2
            vehicle_percentages_others = vehicle_percentages[vehicle_percentages < other_threshold].sum()
            vehicle_percentages = vehicle_percentages[vehicle_percentages >= other_threshold]
            vehicle_percentages['Sonstiges'] = vehicle_percentages_others

            # Kreisdiagramm für Fahrzeugtypen erstellen
            fig, ax = plt.subplots(figsize=(8, 6))
            ax.pie(vehicle_percentages, labels=vehicle_percentages.index, autopct='%1.1f%%', startangle=140, colors=plt.cm.Blues(np.linspace(0.2, 0.8, len(vehicle_percentages))))
            ax.set_title('Verteilung der Fahrzeugtypen (%)', color=primary_color)
            plt.axis('equal')  # Kreisdiagramm als Kreis darstellen
            plt.show()

            # Schritt 2: Zuordnung der Fahrzeugtypen zu OEMs
            brand_mapping = {
                'Skoda Enyaq IV': 'VW', 'VW ID.Buzz': 'VW', 'Cupra Born': 'VW', 'ID4 MY24': 'VW',
                'BMW iX': 'BMW', 'Aiways U5': 'Aiways', 'Aiways U6': 'Aiways', 'VW ID.4': 'VW',
                'ID.5 GTX': 'VW', 'Tesla': 'Tesla', 'VW e-Crafter': 'VW', 'Ford E-Transit': 'Ford',
                'Ford Mustang Mach-E': 'Ford', 'Tesla Model Y': 'Tesla', 'BMW iX1': 'BMW', 'VW e-Golf': 'VW',
                'BMW i3': 'BMW', 'VW eUP': 'VW', 'Seat Mii': 'VW', 'MG5': 'MG', 'Volvo EX30': 'Volvo',
                'smart #1': 'smart', 'smart #3': 'smart', 'Tesla Model 3': 'Tesla', 'Audi E-tron': 'Audi',
                'BMW I4': 'BMW', 'Mercedes EQC': 'Mercedes', 'Hyundai Ioniq 5': 'Hyundai', 'Ioniq 6': 'Hyundai',
                'EV6': 'Kia', 'Porsche Taycan': 'Porsche', 'Audi E-Tron GT': 'Audi', 'Opel eCorsa': 'Opel',
                'Dacia Spring': 'Dacia', 'MG MG4': 'MG', 'Fiat 500e': 'Fiat', 'Peugeot E-208': 'Peugeot',
                'Opel e-Vivaro': 'Opel', 'Mercedes EQV': 'Mercedes', 'Tesla Model S': 'Tesla', 'Audi Q8 e-tron': 'Audi',
                'BMW iX3': 'BMW', 'Hyundai Ioniq': 'Hyundai', 'Mazda MX-30': 'Mazda', 'VW ID.3': 'VW',
                'Renault Zoe': 'Renault', 'Porsche Macan': 'Porsche', 'Mini SE': 'Mini', 'Kia e-Niro': 'Kia',
                'Polestar 2': 'Polestar', 'Volvo XC40 Recharge': 'Volvo', 'Jaguar I-Pace': 'Jaguar'
            }

            # Funktion zur Zuordnung der Fahrzeugtypen zu den OEMs
            def map_to_brand(car_type):
                for model, brand in brand_mapping.items():
                    if model in car_type:
                        return brand
                return 'Sonstiges'

            # Zuordnung der Fahrzeugtypen zu den OEMs
            filtered_df['Brand'] = filtered_df['Car'].dropna().apply(map_to_brand)

            # OEMs zählen und prozentuale Verteilung berechnen
            brand_counts = filtered_df['Brand'].value_counts()
            brand_percentages = brand_counts / brand_counts.sum() * 100

            # OEMs unter 2% als 'Sonstiges' zusammenfassen
            brand_percentages_others = brand_percentages[brand_percentages < other_threshold].sum()
            brand_percentages = brand_percentages[brand_percentages >= other_threshold]
            brand_percentages['Sonstiges'] = brand_percentages_others

            # Kreisdiagramm für OEMs erstellen
            fig, ax = plt.subplots(figsize=(8, 6))
            ax.pie(brand_percentages, labels=brand_percentages.index, autopct='%1.1f%%', startangle=140, colors=plt.cm.Oranges(np.linspace(0.2, 0.8, len(brand_percentages))))
            ax.set_title('Verteilung der OEMs (%)', color=primary_color)
            plt.axis('equal')  # Kreisdiagramm als Kreis darstellen
            plt.show()

    # Eingabefeld für die Seriennummer
    seriennummer_input = widgets.Text(
        value='',
        placeholder='Geben Sie die Seriennummer ein (optional)',
        description='Seriennummer:',
        style={'description_width': 'initial'},
        layout=widgets.Layout(width='300px')
    )

    # Datepicker für die Datumsfilterung
    date_picker_start = widgets.DatePicker(description='Startdatum:')
    date_picker_end = widgets.DatePicker(description='Enddatum:')

    # Button zum Filtern der Daten
    filter_button = widgets.Button(description="Filtern und Visualisieren", button_style='success',
                                   style={'button_color': accent_color})

    def on_filter_click(b):
        seriennummer = seriennummer_input.value
        start_date = date_picker_start.value
        end_date = date_picker_end.value
        filtered_df = filter_data(df, seriennummer=seriennummer, start_date=start_date, end_date=end_date)
        if filtered_df.empty:
            with container:
                container.clear_output()
                display("Keine Daten im angegebenen Zeitraum verfügbar.")
        else:
            plot_vehicle_analysis(filtered_df)

    filter_button.on_click(on_filter_click)

    # Widgets anzeigen
    display(seriennummer_input, date_picker_start, date_picker_end, filter_button, container)
#Datenanahme full control
def show_data_input():
    container = widgets.Output()
    
    # Hinweis für den Benutzer
    display(widgets.HTML("<h3>Data Input</h3>"))
    display(widgets.HTML("Bitte laden Sie Ihre CSV-Datei hoch:"))
    
    # FileUpload-Widget
    upload_widget = widgets.FileUpload(
        accept='.csv',  # Nur CSV-Dateien
        multiple=False
    )
    
    def on_upload_change(change):
        with container:
            for uploaded_filename in upload_widget.value:
                content = upload_widget.value[uploaded_filename]['content']
                df_new = pd.read_csv(
                    io.BytesIO(content),
                    sep=';',  # Passen Sie das Trennzeichen an
                    encoding='ISO-8859-1'
                )
                # Hier können Sie festlegen, wie die hochgeladenen Daten verarbeitet werden sollen
                # Zum Beispiel können Sie sie zu einem globalen DataFrame hinzufügen
                global df_global
                df_global = pd.concat([df_global, df_new], ignore_index=True)
                display('Daten erfolgreich hochgeladen und hinzugefügt.')
                display('Aktuelle Spalten im Datensatz:', df_global.columns.tolist())
    
    upload_widget.observe(on_upload_change, names='value')
    
    display(upload_widget, container)

def show_full_control_dashboard():
    container = widgets.Output()

    # Tabs für die verschiedenen Bereiche
    tab_contents = ['Data Input', 'Presumptions', 'Simulation']
    children = []

    # Funktionen für jeden Tab
    for content in tab_contents:
        out = widgets.Output()
        with out:
            if content == 'Data Input':
                show_data_input()
            elif content == 'Presumptions':
                show_presumptions()
            elif content == 'Simulation':
                show_simulation()
        children.append(out)

    tabs = widgets.Tab()
    tabs.children = children
    for i in range(len(tab_contents)):
        tabs.set_title(i, tab_contents[i])

    display(tabs)

#Annahmen Full control
def show_presumptions():
    container = widgets.Output()
    
    # Hinweis für den Benutzer
    display(widgets.HTML("<h3>Presumptions</h3>"))
    
    # Eingabefelder für die Annahmen
    e_auto_anteil_input = widgets.FloatSlider(
        value=10.0,
        min=0,
        max=100,
        step=0.1,
        description='E-Auto Anteil (%):',
        style={'description_width': 'initial'},
        layout=widgets.Layout(width='400px')
    )
    
    wachstumsrate_e_autos_input = widgets.FloatSlider(
        value=5.0,
        min=0,
        max=100,
        step=0.1,
        description='Wachstumsrate E-Autos (%):',
        style={'description_width': 'initial'},
        layout=widgets.Layout(width='400px')
    )
    
    # Weitere relevante Parameter können hier hinzugefügt werden
    # Zum Beispiel:
    # ladeleistung_input = widgets.FloatText(
    #     value=11.0,
    #     description='Durchschnittliche Ladeleistung (kW):',
    #     style={'description_width': 'initial'},
    #     layout=widgets.Layout(width='400px')
    # )
    
    # Button zum Speichern der Annahmen
    save_button = widgets.Button(
        description="Annahmen speichern",
        button_style='success',
        layout=widgets.Layout(width='200px')
    )
    
    # Output für Bestätigung
    confirmation_output = widgets.Output()
    
    def save_presumptions(b):
        with confirmation_output:
            confirmation_output.clear_output()
            # Hier speichern wir die Annahmen in einer globalen Variablen oder Datei
            # Zum Beispiel in einem Dictionary
            global presumptions
            presumptions = {
                'E-Auto Anteil (%)': e_auto_anteil_input.value,
                'Wachstumsrate E-Autos (%)': wachstumsrate_e_autos_input.value,
                # 'Durchschnittliche Ladeleistung (kW)': ladeleistung_input.value
                # Weitere Parameter hinzufügen
            }
            display("Annahmen wurden erfolgreich gespeichert.")
            display(presumptions)
    
    save_button.on_click(save_presumptions)
    
    # Widgets anzeigen
    display(
        e_auto_anteil_input,
        wachstumsrate_e_autos_input,
        # ladeleistung_input,
        save_button,
        confirmation_output
    )

#BiDI Analyse Full Control:
def bidi_analyse():
    container = widgets.Output()
    display(widgets.HTML("<h4>BiDi Analyse</h4>"))
    
    # Hier können Sie die Logik für die BiDi Analyse implementieren
    # Beispiel: Nutzung der gespeicherten Annahmen
    if 'presumptions' in globals():
        e_auto_anteil = presumptions.get('E-Auto Anteil (%)', 0)
        # Beispielhafte Berechnung
        result = e_auto_anteil * 2  # Dummy-Berechnung
        display(f"Ergebnis der BiDi Analyse basierend auf E-Auto Anteil ({e_auto_anteil}%): {result}")
    else:
        display("Bitte geben Sie zuerst die Annahmen unter 'Presumptions' ein.")
    
    display(container)

#MCS Analse
def mcs_analyse():
    container = widgets.Output()
    display(widgets.HTML("<h4>MCS Analyse</h4>"))
    
    # Hier können Sie die Logik für die MCS Analyse implementieren
    # Beispiel: Nutzung der gespeicherten Annahmen
    if 'presumptions' in globals():
        wachstumsrate = presumptions.get('Wachstumsrate E-Autos (%)', 0)
        # Beispielhafte Berechnung
        result = wachstumsrate * 3  # Dummy-Berechnung
        display(f"Ergebnis der MCS Analyse basierend auf Wachstumsrate ({wachstumsrate}%): {result}")
    else:
        display("Bitte geben Sie zuerst die Annahmen unter 'Presumptions' ein.")
    
    display(container)

#OEM Analyse
def oem_analyse():
    container = widgets.Output()
    display(widgets.HTML("<h4>OEM Analyse</h4>"))
    
    # Hier können Sie die Logik für die OEM Analyse implementieren
    # Beispiel: Analyse basierend auf hochgeladenen Daten
    if not df_global.empty:
        # Beispielhafte Darstellung der OEMs
        if 'Brand' in df_global.columns:
            brand_counts = df_global['Brand'].value_counts()
            fig = px.bar(
                brand_counts,
                x=brand_counts.index,
                y=brand_counts.values,
                labels={'x': 'OEM', 'y': 'Anzahl'},
                title='Anzahl der Fahrzeuge pro OEM'
            )
            fig.update_layout(
                plot_bgcolor=secondary_color,
                paper_bgcolor=secondary_color,
                font_color=primary_color,
                title_font_color=primary_color,
                xaxis_title='OEM',
                yaxis_title='Anzahl',
                title_x=0.5,
                font=dict(size=14),
            )
            fig.update_traces(marker_color=accent_color)
            fig.show()
        else:
            display("Die Spalte 'Brand' ist nicht im Datensatz vorhanden.")
    else:
        display("Keine Daten verfügbar. Bitte laden Sie Daten unter 'Data Input' hoch.")
    
    display(container)

def show_simulation():
    container = widgets.Output()
    
    # Subkarten für Simulation
    subcards = ['BiDi Analyse', 'MCS Analyse', 'OEM Analyse']
    
    # Dropdown für Subkarten
    dropdown_subcard = widgets.Dropdown(
        options=subcards,
        value=subcards[0],
        description='Analyse:',
        style={'description_width': 'initial'},
        layout=widgets.Layout(width='300px')
    )
    
    display(dropdown_subcard, container)
    
    def on_subcard_change(change):
        with container:
            container.clear_output()
            if dropdown_subcard.value == 'BiDi Analyse':
                bidi_analyse()
            elif dropdown_subcard.value == 'MCS Analyse':
                mcs_analyse()
            elif dropdown_subcard.value == 'OEM Analyse':
                oem_analyse()
            else:
                display(f"Subkarte '{dropdown_subcard.value}' wird bald hinzugefügt.")
    
    dropdown_subcard.observe(on_subcard_change, names='value')
    
    # Initiale Ansicht anzeigen
    on_subcard_change(None)


def filtered_documents():
    container = widgets.Output()

    # Mögliche Filteroptionen
    document_types = ['Herstellerdokument', 'MOON Dokument', 'Datenblatt', 'Handbuch', 'Verkaufsunterlage']

    # Checkboxen für die Filter
    checkboxes = [widgets.Checkbox(value=False, description=doc_type) for doc_type in document_types]

    # Button zum Anwenden der Filter
    apply_button = widgets.Button(description="Filter anwenden", button_style='success', style={'button_color': accent_color})

    # Ausgabecontainer
    output = widgets.Output()

    def apply_filters(b):
        with output:
            output.clear_output()
            # Ausgewählte Filter sammeln
            selected_filters = [cb.description for cb in checkboxes if cb.value]
            if not selected_filters:
                display("Bitte wählen Sie mindestens einen Filter aus.")
                return

            # Dokumente filtern
            filtered_docs = documents_df[documents_df['Typ'].isin(selected_filters)]

            if filtered_docs.empty:
                display("Keine Dokumente gefunden.")
            else:
                # Anzeige der gefilterten Dokumente
                for idx, row in filtered_docs.iterrows():
                    display(widgets.HTML(f"<a href='{row['Pfad']}' target='_blank'>{row['Name']}</a>"))

    apply_button.on_click(apply_filters)

    display(widgets.VBox(checkboxes + [apply_button, output]))

def help_guide():
    container = widgets.Output()

    # Beispielhafter Fragebogen
    questions = [
        {
            'question': 'Welche Art von Dokument suchen Sie?',
            'options': ['Herstellerdokument', 'MOON Dokument', 'Datenblatt', 'Handbuch', 'Verkaufsunterlage']
        },
        {
            'question': 'Für welches Produkt benötigen Sie das Dokument?',
            'options': ['HYC50', 'HYC200', 'HYC400', 'C7', 'C8']
        }
    ]

    answers = {}

    def create_question(idx):
        if idx >= len(questions):
            # Dokument basierend auf den Antworten suchen
            with container:
                container.clear_output()
                # Suche nach passenden Dokumenten
                doc_type = answers.get(0)
                product = answers.get(1)
                matched_docs = documents_df[
                    (documents_df['Typ'] == doc_type) & (documents_df['Name'].str.contains(product))
                ]
                if matched_docs.empty:
                    display("Keine passenden Dokumente gefunden.")
                else:
                    for idx, row in matched_docs.iterrows():
                        display(widgets.HTML(f"<a href='{row['Pfad']}' target='_blank'>{row['Name']}</a>"))
            return

        question = questions[idx]
        options = question['options']
        radio_buttons = widgets.RadioButtons(
            options=options,
            description=question['question'],
            style={'description_width': 'initial'}
        )

        next_button = widgets.Button(description="Weiter", button_style='success', style={'button_color': accent_color})

        def on_next(b):
            answers[idx] = radio_buttons.value
            container.clear_output()
            create_question(idx + 1)

        next_button.on_click(on_next)

        with container:
            container.clear_output()
            display(radio_buttons, next_button)

    # Start mit der ersten Frage
    create_question(0)

    display(container)
def full_list_documents():
    container = widgets.Output()

    # Suchleiste
    search_input = widgets.Text(
        value='',
        placeholder='Dokumentname eingeben',
        description='Suche:',
        style={'description_width': 'initial'},
        layout=widgets.Layout(width='400px')
    )

    # Ausgabecontainer
    output = widgets.Output()

    def on_search_change(change):
        with output:
            output.clear_output()
            search_term = search_input.value.lower()
            if not search_term:
                # Alle Dokumente anzeigen
                docs_to_display = documents_df
            else:
                # Dokumente filtern
                docs_to_display = documents_df[documents_df['Name'].str.lower().str.contains(search_term)]
            if docs_to_display.empty:
                display("Keine Dokumente gefunden.")
            else:
                for idx, row in docs_to_display.iterrows():
                    display(widgets.HTML(f"<a href='{row['Pfad']}' target='_blank'>{row['Name']}</a>"))

    search_input.observe(on_search_change, names='value')

    display(search_input, output)
    # Initiale Anzeige aller Dokumente
    on_search_change(None)
def show_documents_dashboard():
    container = widgets.Output()

    # Subkarten für Documents
    subkarten = ['Filtered Document', 'Help Guide', 'Full List']

    # Dropdown für Subkarten
    dropdown_subkarte = widgets.Dropdown(
        options=subkarten,
        value=subkarten[0],
        description='Option:',
        style={'description_width': 'initial'},
        layout=widgets.Layout(width='300px')
    )

    display(dropdown_subkarte, container)

    def on_subkarte_change(change):
        with container:
            container.clear_output()
            if dropdown_subkarte.value == 'Filtered Document':
                filtered_documents()
            elif dropdown_subkarte.value == 'Help Guide':
                help_guide()
            elif dropdown_subkarte.value == 'Full List':
                full_list_documents()
            else:
                display(f"Subkarte '{dropdown_subkarte.value}' wird bald hinzugefügt.")

    dropdown_subkarte.observe(on_subkarte_change, names='value')

    # Initiales Dashboard anzeigen
    on_subkarte_change(None)

# Vertriebs-Dashboard: Chargerausnutzung
def vertrieb_chargerausnutzung(df):
    container = widgets.Output()

    def plot_data(filtered_df, seriennummer=None):
        with container:
            container.clear_output()

            # Überprüfen, ob 'Start Time' in den Spalten vorhanden ist
            if 'Start Time' not in filtered_df.columns:
                display("Die Spalte 'Start Time' ist nicht im DataFrame vorhanden.")
                return

            # Daten konvertieren
            filtered_df['Start Time'] = pd.to_datetime(filtered_df['Start Time'], errors='coerce', dayfirst=True)
            filtered_df['Weekday'] = filtered_df['Start Time'].dt.day_name()
            filtered_df['Hour'] = filtered_df['Start Time'].dt.hour

            weekday_order = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
            sessions_per_weekday = filtered_df['Weekday'].value_counts().reindex(weekday_order)
            average_sessions_per_weekday = sessions_per_weekday / filtered_df['Start Time'].dt.date.nunique()

            # Plot: Durchschnittliche Anzahl der Ladevorgänge pro Wochentag
            title = 'Durchschnittliche Anzahl der Ladevorgänge pro Wochentag'
            if seriennummer:
                title += f' für Seriennummer {seriennummer}'

            # CI-Anpassungen für die Grafik
            fig = px.bar(
                average_sessions_per_weekday, 
                x=average_sessions_per_weekday.index, 
                y=average_sessions_per_weekday.values, 
                labels={'x': 'Wochentag', 'y': 'Durchschnittliche Anzahl der Ladevorgänge'},
                title=title
            )
            fig.update_layout(
                plot_bgcolor=secondary_color,
                paper_bgcolor=secondary_color,
                font_color=primary_color,
                title_font_color=primary_color,
                xaxis_title='Wochentag',
                yaxis_title='Durchschnittliche Anzahl der Ladevorgänge',
                title_x=0.5, 
                xaxis_tickangle=-45,
                bargap=0.2,
                yaxis=dict(
                    showgrid=True,
                    gridcolor=primary_color,
                    linecolor=primary_color
                ),
                xaxis=dict(
                    showgrid=False,
                    linecolor=primary_color
                ),
                title_font=dict(size=18),
                font=dict(size=14)
            )
            fig.update_traces(marker_color=accent_color)

            fig.show()

            # Plot: Anzahl der Ladevorgänge pro Stunde des Tages
            sessions_per_hour = filtered_df['Hour'].value_counts().sort_index()
            fig_hour = px.bar(
                sessions_per_hour, 
                x=sessions_per_hour.index, 
                y=sessions_per_hour.values, 
                labels={'x': 'Stunde des Tages', 'y': 'Anzahl der Ladevorgänge'},
                title=f'Anzahl der Ladevorgänge pro Stunde des Tages' + (f' für Seriennummer {seriennummer}' if seriennummer else '')
            )
            fig_hour.update_layout(
                plot_bgcolor=secondary_color,
                paper_bgcolor=secondary_color,
                font_color=primary_color,
                title_font_color=primary_color,
                xaxis_title='Stunde des Tages',
                yaxis_title='Anzahl der Ladevorgänge',
                title_x=0.5, 
                bargap=0.2,
                yaxis=dict(
                    showgrid=True,
                    gridcolor=primary_color,
                    linecolor=primary_color
                ),
                xaxis=dict(
                    showgrid=False,
                    linecolor=primary_color
                ),
                title_font=dict(size=18),
                font=dict(size=14)
            )
            fig_hour.update_traces(marker_color=accent_color)

            fig_hour.show()

    # Eingabefeld für die Seriennummer
    seriennummer_input = widgets.Text(
        value='',
        placeholder='Geben Sie die Seriennummer ein (optional)',
        description='Seriennummer:',
        style={'description_width': 'initial'},
        layout=widgets.Layout(width='300px')
    )

    # Datepicker für die Datumsfilterung
    date_picker_start = widgets.DatePicker(
        description='Startdatum:',
        disabled=False
    )
    date_picker_end = widgets.DatePicker(
        description='Enddatum:',
        disabled=False
    )

    # Button zum Filtern der Daten basierend auf der Seriennummer und dem Datum
    filter_button = widgets.Button(description="Filtern und Visualisieren", button_style='success',
                                   style={'button_color': accent_color},
                                   layout=widgets.Layout(width='200px'))
    # Reset-Button zum Zurücksetzen der Filter
    reset_button = widgets.Button(description="Reset", button_style='warning',
                                  layout=widgets.Layout(width='200px'))

    def on_filter_click(b):
        seriennummer = seriennummer_input.value
        start_date = date_picker_start.value
        end_date = date_picker_end.value

        # Daten nach Seriennummer und Zeit filtern
        filtered_df = filter_data(df, seriennummer, start_date, end_date)

        if filtered_df.empty:
            with container:
                container.clear_output()
                display(f"Keine Daten für Seriennummer: {seriennummer} im angegebenen Zeitraum.")
        else:
            plot_data(filtered_df, seriennummer)

    def on_reset_click(b):
        seriennummer_input.value = ''
        date_picker_start.value = None
        date_picker_end.value = None
        plot_data(df)

    filter_button.on_click(on_filter_click)
    reset_button.on_click(on_reset_click)

    # Widgets anzeigen
    display(seriennummer_input, date_picker_start, date_picker_end, filter_button, reset_button, container)


# Liste von Postleitzahlen und ihren entsprechenden Städten
plz_to_city = {
    "1010": "Wien", "1020": "Wien", "1030": "Wien",
    "2000": "Stockerau", "3100": "St. Pölten", "5020": "Salzburg",
    # Weitere Einträge hinzufügen...
}

# Amortisationsrechnung-Dashboard
def amortisationsrechnung_dashboard():
    container = widgets.Output()

    # Eingabefelder für die Projektinformationen
    postleitzahl_input = widgets.Text(
        value='',
        placeholder='Postleitzahl eingeben',
        description='Postleitzahl:',
        style={'description_width': 'initial'},
        layout=widgets.Layout(width='300px')
    )

    hyc50_input = widgets.IntText(
        value=0,
        description='HYC50:',
        style={'description_width': 'initial'},
        layout=widgets.Layout(width='150px')
    )

    hyc200_input = widgets.IntText(
        value=0,
        description='HYC200:',
        style={'description_width': 'initial'},
        layout=widgets.Layout(width='150px')
    )

    hyc400_input = widgets.IntText(
        value=0,
        description='HYC400:',
        style={'description_width': 'initial'},
        layout=widgets.Layout(width='150px')
    )

    c7_input = widgets.IntText(
        value=0,
        description='C7:',
        style={'description_width': 'initial'},
        layout=widgets.Layout(width='150px')
    )

    c8_input = widgets.IntText(
        value=0,
        description='C8:',
        style={'description_width': 'initial'},
        layout=widgets.Layout(width='150px')
    )

    kosten_input = widgets.FloatText(
        value=0.0,
        description='Gesamtkosten (EUR):',
        style={'description_width': 'initial'},
        layout=widgets.Layout(width='300px')
    )

    miete_input = widgets.FloatText(
        value=0.0,
        description='Miete pro Jahr (EUR):',
        style={'description_width': 'initial'},
        layout=widgets.Layout(width='300px')
    )

    einkaufspreis_input = widgets.FloatText(
        value=0.0,
        description='Einkaufspreis Strom (EUR/kWh):',
        style={'description_width': 'initial'},
        layout=widgets.Layout(width='300px')
    )

    verkaufspreis_input = widgets.FloatText(
        value=0.0,
        description='Verkaufspreis Strom (EUR/kWh):',
        style={'description_width': 'initial'},
        layout=widgets.Layout(width='300px')
    )

    thg_input = widgets.FloatText(
        value=0.0,
        description='THG-Quote (EUR/kWh):',
        style={'description_width': 'initial'},
        layout=widgets.Layout(width='300px')
    )

    # Button zum Berechnen der Amortisation
    berechnen_button = widgets.Button(
        description="Amortisation berechnen",
        button_style='success',
        layout=widgets.Layout(width='200px')
    )

    # Output für die Amortisationsrechnung
    ergebnis_output = widgets.Output()

    def berechne_amortisation(b):
        with ergebnis_output:
            ergebnis_output.clear_output()
            # Hier wird die Berechnung der Amortisationszeit durchgeführt
            # Beispiel: einfache Berechnung der Rentabilität
            try:
                gesamt_kosten = kosten_input.value
                miete_pro_jahr = miete_input.value
                verkaufspreis = verkaufspreis_input.value
                einkaufspreis = einkaufspreis_input.value
                thg_quote = thg_input.value

                # Berechnung der Einnahmen durch Verkauf von Strom + THG-Quote
                einnahmen_pro_kwh = verkaufspreis + thg_quote
                gewinn_pro_kwh = einnahmen_pro_kwh - einkaufspreis

                # Beispielhafte Schätzung der Ladevolumina durch Anzahl der Stationen (HYC, C7, C8)
                anzahl_hyc = hyc50_input.value + hyc200_input.value + hyc400_input.value
                anzahl_c = c7_input.value + c8_input.value
                volumen_pro_station = 50000  # Annahme: jede Station liefert 50.000 kWh pro Jahr

                gesamt_volumen = (anzahl_hyc + anzahl_c) * volumen_pro_station
                gewinn_pro_jahr = gesamt_volumen * gewinn_pro_kwh

                # Amortisationszeit berechnen
                amortisationszeit = gesamt_kosten / (gewinn_pro_jahr - miete_pro_jahr)
                display(f"Amortisationszeit: {amortisationszeit:.2f} Jahre")
            except ZeroDivisionError:
                display("Fehler: Division durch Null. Überprüfen Sie Ihre Eingaben.")
            except Exception as e:
                display(f"Fehler bei der Berechnung: {e}")

    berechnen_button.on_click(berechne_amortisation)

    # Widgets anzeigen
    display(
        postleitzahl_input, 
        hyc50_input, hyc200_input, hyc400_input, 
        c7_input, c8_input, 
        kosten_input, miete_input, 
        einkaufspreis_input, verkaufspreis_input, thg_input, 
        berechnen_button, ergebnis_output
    )


# Vertriebs-Dashboard: Chargeranalyse (mit Heatmap zwischen Tag und Stunde und Seriennummer-Auswahl)
def vertrieb_chargeranalyse(df):
    container = widgets.Output()

    def plot_heatmap(filtered_df, seriennummer=None):
        with container:
            container.clear_output()
            if 'Charger Serial Number' not in filtered_df.columns or 'Start Time' not in filtered_df.columns:
                display("Die erforderlichen Spalten sind nicht im DataFrame vorhanden.")
                return

            # Daten vorbereiten
            filtered_df['Date'] = filtered_df['Start Time'].dt.date
            filtered_df['Hour'] = filtered_df['Start Time'].dt.hour

            charger_usage = filtered_df.groupby(['Hour', 'Date']).size().reset_index(name='Sessions')
            pivot_df = charger_usage.pivot(index='Hour', columns='Date', values='Sessions').fillna(0)

            # Heatmap erstellen
            fig = px.imshow(
                pivot_df,
                labels=dict(x="Datum", y="Stunde", color="Anzahl der Sessions"),
                aspect="auto",
                title="Heatmap der Charger-Nutzung (Stunde vs. Datum)" + (f' für Seriennummer {seriennummer}' if seriennummer else '')
            )

            # Layout-Anpassungen
            fig.update_layout(
                plot_bgcolor=secondary_color,
                paper_bgcolor=secondary_color,
                font_color=primary_color,
                title_font_color=primary_color,
                xaxis_title='Datum',
                yaxis_title='Stunde',
                title_x=0.5,
                font=dict(size=14),
            )

            fig.show()

    # Eingabefeld für die Seriennummer
    seriennummer_input = widgets.Text(
        value='',
        placeholder='Geben Sie die Seriennummer ein (optional)',
        description='Seriennummer:',
        style={'description_width': 'initial'},
        layout=widgets.Layout(width='300px')
    )

    # Datepicker für die Datumsfilterung
    date_picker_start = widgets.DatePicker(description='Startdatum:')
    date_picker_end = widgets.DatePicker(description='Enddatum:')

    # Button zum Filtern der Daten
    filter_button = widgets.Button(description="Filtern und Visualisieren", button_style='success',
                                   style={'button_color': accent_color})

    def on_filter_click(b):
        seriennummer = seriennummer_input.value
        start_date = date_picker_start.value
        end_date = date_picker_end.value
        filtered_df = filter_data(df, seriennummer=seriennummer, start_date=start_date, end_date=end_date)
        if filtered_df.empty:
            with container:
                container.clear_output()
                display("Keine Daten im angegebenen Zeitraum verfügbar.")
        else:
            plot_heatmap(filtered_df, seriennummer)

    filter_button.on_click(on_filter_click)

    # Widgets anzeigen
    display(seriennummer_input, date_picker_start, date_picker_end, filter_button, container)

# After Sales Dashboard: Fehlerquote und Fehlerarten
def after_sales_dashboard():
    df = df_global

    # Subkategorien für After Sales
    subkarten = ['Fehlerquote', 'Fehlerarten', 'Charger Fehleranalyse']

    # Dropdown für Subkarten
    dropdown_subkarte = widgets.Dropdown(
        options=subkarten,
        value=subkarten[0],  # Standardwert
        description='Subkarte:',
        style={'description_width': 'initial'},
        layout=widgets.Layout(width='300px')
    )

    container = widgets.Output()
    display(dropdown_subkarte, container)

    def on_subkarte_change(change):
        with container:
            container.clear_output()
            if dropdown_subkarte.value == 'Fehlerquote':
                fehlerquote_dashboard(df)
            elif dropdown_subkarte.value == 'Fehlerarten':
                fehlerarten_dashboard(df)
            else:
                display(f"Subkarte '{dropdown_subkarte.value}' wird bald hinzugefügt.")

    dropdown_subkarte.observe(on_subkarte_change, names='value')

    # Initiales Dashboard anzeigen
    on_subkarte_change(None)

def fehlerquote_dashboard(df):
    container = widgets.Output()

    def plot_data(filtered_df, seriennummer=None):
        with container:
            container.clear_output()

            if 'HYC Errorcode' not in filtered_df.columns:
                display("Die Spalte 'HYC Errorcode' ist nicht im DataFrame vorhanden.")
                return

            total_sessions = len(filtered_df)
            error_sessions = filtered_df['HYC Errorcode'].notna().sum()
            error_rate = (error_sessions / total_sessions) * 100

            # Visualisierung der Fehlerquote mit CI-Vorgaben
            labels = ['Fehlerfrei', 'Fehlerhaft']
            values = [total_sessions - error_sessions, error_sessions]

            fig = px.pie(
                names=labels,
                values=values,
                title='Fehlerquote' + (f' für Seriennummer {seriennummer}' if seriennummer else ''),
                hole=0.4,
                color_discrete_sequence=[accent_color, primary_color]
            )

            fig.update_layout(
                plot_bgcolor=secondary_color,
                paper_bgcolor=secondary_color,
                font_color=primary_color,
                title_font_color=primary_color,
                title_x=0.5,
                font=dict(size=14),
            )

            fig.show()

    # Eingabefeld für die Seriennummer
    seriennummer_input = widgets.Text(
        value='',
        placeholder='Geben Sie die Seriennummer ein (optional)',
        description='Seriennummer:',
        style={'description_width': 'initial'},
        layout=widgets.Layout(width='300px')
    )

    # Datepicker für die Datumsfilterung
    date_picker_start = widgets.DatePicker(description='Startdatum:')
    date_picker_end = widgets.DatePicker(description='Enddatum:')

    # Button zum Filtern der Daten
    filter_button = widgets.Button(description="Filtern und Visualisieren", button_style='success',
                                   style={'button_color': accent_color})

    def on_filter_click(b):
        seriennummer = seriennummer_input.value
        start_date = date_picker_start.value
        end_date = date_picker_end.value
        filtered_df = filter_data(df, seriennummer=seriennummer, start_date=start_date, end_date=end_date)
        if filtered_df.empty:
            with container:
                container.clear_output()
                display("Keine Daten im angegebenen Zeitraum verfügbar.")
        else:
            plot_data(filtered_df, seriennummer)

    filter_button.on_click(on_filter_click)

    # Widgets anzeigen
    display(seriennummer_input, date_picker_start, date_picker_end, filter_button, container)

def fehlerarten_dashboard(df):
    container = widgets.Output()

    def plot_data(filtered_df):
        with container:
            container.clear_output()

            if 'HYC Errorcode' not in filtered_df.columns:
                display("Die Spalte 'HYC Errorcode' ist nicht im DataFrame vorhanden.")
                return

            error_types = filtered_df['HYC Errorcode'].dropna().value_counts()

            # Plot: Balkendiagramm der Fehlerarten mit CI-Vorgaben
            fig = px.bar(
                error_types,
                x=error_types.index,
                y=error_types.values,
                labels={'x': 'Fehlercode', 'y': 'Anzahl'},
                title='Häufigkeit der Fehlerarten'
            )

            # Layout-Anpassungen
            fig.update_layout(
                plot_bgcolor=secondary_color,
                paper_bgcolor=secondary_color,
                font_color=primary_color,
                title_font_color=primary_color,
                xaxis_title='Fehlercode',
                yaxis_title='Anzahl',
                title_x=0.5,
                font=dict(size=14),
            )
            fig.update_traces(marker_color=accent_color)

            fig.show()

    # Eingabefeld für die Seriennummer
    seriennummer_input = widgets.Text(
        value='',
        placeholder='Geben Sie die Seriennummer ein (optional)',
        description='Seriennummer:',
        style={'description_width': 'initial'},
        layout=widgets.Layout(width='300px')
    )

    # Datepicker für die Datumsfilterung
    date_picker_start = widgets.DatePicker(description='Startdatum:')
    date_picker_end = widgets.DatePicker(description='Enddatum:')

    # Button zum Filtern der Daten
    filter_button = widgets.Button(description="Filtern und Visualisieren", button_style='success',
                                   style={'button_color': accent_color})

    def on_filter_click(b):
        seriennummer = seriennummer_input.value
        start_date = date_picker_start.value
        end_date = date_picker_end.value
        filtered_df = filter_data(df, seriennummer=seriennummer, start_date=start_date, end_date=end_date)
        if filtered_df.empty:
            with container:
                container.clear_output()
                display("Keine Daten im angegebenen Zeitraum verfügbar.")
        else:
            plot_data(filtered_df)

    filter_button.on_click(on_filter_click)

    # Widgets anzeigen
    display(seriennummer_input, date_picker_start, date_picker_end, filter_button, container)

# Datenpflege-Dashboard mit CSV-Upload
def datenpflege_dashboard():
    container = widgets.Output()

    # Optionen für Datenpflege
    options = ['Daten hochladen', 'Daten anzeigen']

    # Dropdown-Menü für Optionen
    dropdown_option = widgets.Dropdown(
        options=options,
        value=options[0],
        description='Option:',
        style={'description_width': 'initial'},
    )

    display(dropdown_option, container)

    def on_option_change(change):
        with container:
            container.clear_output()
            if dropdown_option.value == 'Daten hochladen':
                datenpflege_upload()
            elif dropdown_option.value == 'Daten anzeigen':
                datenpflege_anzeigen()

    dropdown_option.observe(on_option_change, names='value')

    # Initiale Ansicht
    on_option_change(None)

def datenpflege_upload():
    container = widgets.Output()
    upload_widget = widgets.FileUpload(
        accept='.csv',  # Nur CSV-Dateien
        multiple=False
    )

    def on_upload_change(change):
        with container:
            for uploaded_filename in upload_widget.value:
                content = upload_widget.value[uploaded_filename]['content']
                df_new = pd.read_csv(
                    io.BytesIO(content),
                    sep=';',  # Passen Sie das Trennzeichen an
                    encoding='ISO-8859-1'
                )
                # Globale df_global aktualisieren
                global df_global
                df_global = pd.concat([df_global, df_new], ignore_index=True)
                display('Daten erfolgreich hochgeladen und hinzugefügt.')
                display('Aktuelle Spalten im Datensatz:', df_global.columns.tolist())

    upload_widget.observe(on_upload_change, names='value')

    display(upload_widget, container)

def datenpflege_anzeigen():
    with container:
        display(df_global.head())

# Reporting-Dashboard mit Seriennummer-Auswahl und Word-Generierung
def reporting_dashboard():
    container = widgets.Output()

    # Eingabefeld für die Seriennummer
    seriennummer_input = widgets.Text(
        value='',
        placeholder='Geben Sie die Seriennummer ein',
        description='Seriennummer:',
        style={'description_width': 'initial'},
        layout=widgets.Layout(width='300px')
    )

    # Button zum Generieren des Berichts
    report_button = widgets.Button(description="Bericht generieren", button_style='success',
                                   style={'button_color': accent_color})

    def generate_report(b):
        seriennummer = seriennummer_input.value
        if not seriennummer:
            with container:
                container.clear_output()
                display("Bitte geben Sie eine Seriennummer ein.")
            return

        filtered_df = filter_data(df_global, seriennummer=seriennummer)
        if filtered_df.empty:
            with container:
                container.clear_output()
                display(f"Keine Daten für Seriennummer: {seriennummer}")
        else:
            # Generieren der Grafiken
            # Grafik 1: Durchschnittliche Anzahl der Ladevorgänge pro Wochentag
            filtered_df['Start Time'] = pd.to_datetime(filtered_df['Start Time'], errors='coerce', dayfirst=True)
            filtered_df['Weekday'] = filtered_df['Start Time'].dt.day_name()
            weekday_order = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
            sessions_per_weekday = filtered_df['Weekday'].value_counts().reindex(weekday_order)
            average_sessions_per_weekday = sessions_per_weekday / filtered_df['Start Time'].dt.date.nunique()
            fig1 = px.bar(
                average_sessions_per_weekday,
                x=average_sessions_per_weekday.index,
                y=average_sessions_per_weekday.values,
                labels={'x': 'Wochentag', 'y': 'Durchschnittliche Anzahl der Ladevorgänge'},
                title='Durchschnittliche Anzahl der Ladevorgänge pro Wochentag'
            )

            # Grafik 2: Anzahl der Ladevorgänge pro Stunde des Tages
            filtered_df['Hour'] = filtered_df['Start Time'].dt.hour
            sessions_per_hour = filtered_df['Hour'].value_counts().sort_index()
            fig2 = px.bar(
                sessions_per_hour,
                x=sessions_per_hour.index,
                y=sessions_per_hour.values,
                labels={'x': 'Stunde des Tages', 'y': 'Anzahl der Ladevorgänge'},
                title='Anzahl der Ladevorgänge pro Stunde des Tages'
            )

            # Word-Dokument erstellen
            doc = Document()
            doc.add_heading(f'Report für Charger {seriennummer}', 0)

            # Grafik 1 hinzufügen
            img_bytes = fig1.to_image(format="png")
            img_stream = io.BytesIO(img_bytes)
            doc.add_heading('Durchschnittliche Anzahl der Ladevorgänge pro Wochentag', level=1)
            doc.add_picture(img_stream, width=Inches(6))

            # Grafik 2 hinzufügen
            img_bytes = fig2.to_image(format="png")
            img_stream = io.BytesIO(img_bytes)
            doc.add_heading('Anzahl der Ladevorgänge pro Stunde des Tages', level=1)
            doc.add_picture(img_stream, width=Inches(6))

            # Speichern des Dokuments
            report_filename = f'Report_Charger_{seriennummer}.docx'
            doc.save(report_filename)

            with container:
                container.clear_output()
                display(f'Bericht "{report_filename}" wurde erfolgreich generiert und gespeichert.')
                # Möglichkeit zum Herunterladen des Berichts
                from IPython.display import FileLink
                display(FileLink(report_filename))

    report_button.on_click(generate_report)

    # Widgets anzeigen
    display(seriennummer_input, report_button, container)

# Marketing-Dashboard mit Europakarte
def marketing_dashboard():
    df = df_global

    # Subkarten für Marketing
    subkarten = ['Kundenanalyse', 'Länderanalyse']

    # Dropdown für Subkarten
    dropdown_subkarte = widgets.Dropdown(
        options=subkarten,
        value='Länderanalyse',
        description='Subkarte:',
        style={'description_width': 'initial'},
        layout=widgets.Layout(width='300px')
    )

    container = widgets.Output()
    display(dropdown_subkarte, container)

    def on_subkarte_change(change):
        with container:
            container.clear_output()
            if dropdown_subkarte.value == 'Länderanalyse':
                laenderanalyse_dashboard(df)
            else:
                display(f"Subkarte '{dropdown_subkarte.value}' wird bald hinzugefügt.")

    dropdown_subkarte.observe(on_subkarte_change, names='value')

    # Initiales Dashboard anzeigen
    on_subkarte_change(None)

def laenderanalyse_dashboard(df):
    container = widgets.Output()

    def plot_map():
        with container:
            container.clear_output()

            # Beispiel-Daten erstellen
            country_counts = pd.DataFrame({
                'Country': ['Austria', 'Germany', 'Romania'],
                'Counts': [50, 100, 70]
            })

            # Europakarte plotten
            fig = px.choropleth(
                country_counts,
                locations='Country',
                locationmode='country names',
                color='Counts',
                scope='europe',
                title='Anzahl der Charger pro Land',
                color_continuous_scale=px.colors.sequential.Blues
            )
            fig.update_layout(
                plot_bgcolor=secondary_color,
                paper_bgcolor=secondary_color,
                font_color=primary_color,
                title_font_color=primary_color,
                title_x=0.5,
                font=dict(size=14),
            )
            fig.show()

    # Karte plotten
    plot_map()

    # Widgets anzeigen
    display(container)

# Produktmanagement-Dashboard mit Nutzeranalyse
# Produktmanagement-Dashboard mit Ländervergleich
def produktmanagement_dashboard(df):
    container = widgets.Output()

    subkarten = ['Produktausnutzung', 'Nutzeranalyse', 'Ländervergleich', 'Performance']

    dropdown_subkarte = widgets.Dropdown(
        options=subkarten,
        value=subkarten[2],  # Standardmäßig wird 'Ländervergleich' ausgewählt
        description='Subkarte:',
        style={'description_width': 'initial'},
    )

    display(dropdown_subkarte, container)

    def on_subkarte_change(change):
        with container:
            container.clear_output()
            if dropdown_subkarte.value == 'Produktausnutzung':
                produktausnutzung_dashboard(df)
            elif dropdown_subkarte.value == 'Nutzeranalyse':
                nutzeranalyse_dashboard(df)
            elif dropdown_subkarte.value == 'Ländervergleich':
                laendervergleich_dashboard(df)  # Aufruf der neuen Funktion
            else:
                display(f"Subkarte '{dropdown_subkarte.value}' wird bald hinzugefügt.")

    dropdown_subkarte.observe(on_subkarte_change, names='value')

    # Initiales Dashboard anzeigen
    on_subkarte_change(None)


    # Initiale Ansicht anzeigen
    on_subkarte_change(None)


    # Initiales Dashboard anzeigen
    on_subkarte_change(None)

def nutzeranalyse_dashboard(df):
    container = widgets.Output()

    def plot_data(filtered_df):
        with container:
            container.clear_output()

            required_columns = ['SoC Start (%)', 'SoC Stop (%)', 'Chargingsession ID', 'Start Time', 'End Time']
            for col in required_columns:
                if col not in filtered_df.columns:
                    display(f"Die erforderliche Spalte '{col}' ist nicht im DataFrame vorhanden.")
                    return

            # Anzahl der Ladevorgänge berechnen
            num_sessions = filtered_df['Chargingsession ID'].nunique()
            display(f"Anzahl der durchgeführten Ladevorgänge: {num_sessions}")

            # Konvertierung der 'Start Time' und 'End Time' Spalten in das Datetime-Format
            filtered_df['Start Time'] = pd.to_datetime(filtered_df['Start Time'], errors='coerce', dayfirst=True)
            filtered_df['End Time'] = pd.to_datetime(filtered_df['End Time'], errors='coerce', dayfirst=True)

            # Berechnung der Ladedauer in Minuten
            filtered_df['Ladedauer (Minuten)'] = (filtered_df['End Time'] - filtered_df['Start Time']).dt.total_seconds() / 60

            # Ausschließen von Ladevorgängen mit negativen oder null Ladedauer
            df_filtered = filtered_df[filtered_df['Ladedauer (Minuten)'] > 0]

            # Optional: Ausschließen von Ladevorgängen mit Ladedauer über 120 Minuten
            df_filtered = df_filtered[df_filtered['Ladedauer (Minuten)'] < 120]

            # Histogramm der Ladedauer
            fig_duration_hist = px.histogram(
                df_filtered,
                x='Ladedauer (Minuten)',
                nbins=20,
                title='Verteilung der Ladedauer (Minuten)',
                labels={'Ladedauer (Minuten)': 'Ladedauer (Minuten)', 'count': 'Anzahl der Ladevorgänge'}
            )
            fig_duration_hist.update_layout(
                plot_bgcolor=secondary_color,
                paper_bgcolor=secondary_color,
                font_color=primary_color,
                title_font_color=primary_color,
                xaxis_title='Ladedauer (Minuten)',
                yaxis_title='Anzahl der Ladevorgänge',
                title_x=0.5,
                font=dict(size=14),
            )
            fig_duration_hist.update_traces(marker_color=accent_color)
            fig_duration_hist.show()

            # Boxplot der Ladedauer
            fig_duration_box = px.box(
                df_filtered,
                x='Ladedauer (Minuten)',
                title='Boxplot der Ladedauer (Minuten)',
                labels={'Ladedauer (Minuten)': 'Ladedauer (Minuten)'}
            )
            fig_duration_box.update_layout(
                plot_bgcolor=secondary_color,
                paper_bgcolor=secondary_color,
                font_color=primary_color,
                title_font_color=primary_color,
                xaxis_title='Ladedauer (Minuten)',
                title_x=0.5,
                font=dict(size=14),
            )
            fig_duration_box.update_traces(marker_color=accent_color)
            fig_duration_box.show()

            # Histogramm der Verteilung der SoC Start (%)
            fig_soc_start = px.histogram(
                filtered_df,
                x='SoC Start (%)',
                nbins=20,
                title='Verteilung der Start SoC (%) der Ladevorgänge',
                labels={'SoC Start (%)': 'Start SoC (%)', 'count': 'Anzahl der Ladevorgänge'}
            )
            fig_soc_start.update_layout(
                plot_bgcolor=secondary_color,
                paper_bgcolor=secondary_color,
                font_color=primary_color,
                title_font_color=primary_color,
                xaxis_title='Start SoC (%)',
                yaxis_title='Anzahl der Ladevorgänge',
                title_x=0.5,
                font=dict(size=14),
            )
            fig_soc_start.update_traces(marker_color=accent_color)
            fig_soc_start.show()

            # Histogramm der Verteilung der SoC Stop (%)
            fig_soc_stop = px.histogram(
                filtered_df,
                x='SoC Stop (%)',
                nbins=20,
                title='Verteilung der End SoC (%) der Ladevorgänge',
                labels={'SoC Stop (%)': 'End SoC (%)', 'count': 'Anzahl der Ladevorgänge'}
            )
            fig_soc_stop.update_layout(
                plot_bgcolor=secondary_color,
                paper_bgcolor=secondary_color,
                font_color=primary_color,
                title_font_color=primary_color,
                xaxis_title='End SoC (%)',
                yaxis_title='Anzahl der Ladevorgänge',
                title_x=0.5,
                font=dict(size=14),
            )
            fig_soc_stop.update_traces(marker_color=accent_color)
            fig_soc_stop.show()

            # --- NEU: Europakarten hinzufügen ---
            def plot_europe_map(filtered_df, title):
                fig = px.choropleth(
                    filtered_df,
                    locations='Country', 
                    locationmode='country names', 
                    color='Sessions',
                    title=title,
                    scope='europe',
                    color_continuous_scale=px.colors.sequential.Blues
                )
                fig.update_layout(
                    plot_bgcolor=secondary_color,
                    paper_bgcolor=secondary_color,
                    font_color=primary_color,
                    title_font_color=primary_color,
                    title_x=0.5,
                    font=dict(size=14),
                )
                fig.show()

            # Beispieldaten für Österreich, Deutschland und Rumänien
            seriennummer_osterreich = '19BZ00401'
            seriennummer_deutschland = '23BZ2234B'
            seriennummer_rumaenien = '20BZ00001'

            # Daten nach Ländern filtern
            filtered_df_osterreich = filter_data(df, seriennummer=seriennummer_osterreich)
            filtered_df_deutschland = filter_data(df, seriennummer=seriennummer_deutschland)
            filtered_df_rumaenien = filter_data(df, seriennummer=seriennummer_rumaenien)

            # Erstellen von fiktiven Daten zu den Sessions für die Karten
            country_data = pd.DataFrame({
                'Country': ['Austria', 'Germany', 'Romania'],
                'Sessions': [
                    filtered_df_osterreich.shape[0],
                    filtered_df_deutschland.shape[0],
                    filtered_df_rumaenien.shape[0]
                ]
            })

            # Titel für die Karten
            title_osterreich = f'Charger Sessions in Österreich (Seriennummer {seriennummer_osterreich})'
            title_deutschland = f'Charger Sessions in Deutschland (Seriennummer {seriennummer_deutschland})'
            title_rumaenien = f'Charger Sessions in Rumänien (Seriennummer {seriennummer_rumaenien})'

            # Interaktive Europakarten anzeigen
            plot_europe_map(country_data, title='Sessions in den Ländern')
            plot_europe_map(pd.DataFrame({'Country': ['Austria'], 'Sessions': [country_data['Sessions'][0]]}), title_osterreich)
            plot_europe_map(pd.DataFrame({'Country': ['Germany'], 'Sessions': [country_data['Sessions'][1]]}), title_deutschland)
            plot_europe_map(pd.DataFrame({'Country': ['Romania'], 'Sessions': [country_data['Sessions'][2]]}), title_rumaenien)

    # Eingabefeld für die Seriennummer
    seriennummer_input = widgets.Text(
        value='',
        placeholder='Geben Sie die Seriennummer ein (optional)',
        description='Seriennummer:',
        style={'description_width': 'initial'},
        layout=widgets.Layout(width='300px')
    )

    # Datepicker für die Datumsfilterung
    date_picker_start = widgets.DatePicker(description='Startdatum:')
    date_picker_end = widgets.DatePicker(description='Enddatum:')

    # Button zum Filtern der Daten
    filter_button = widgets.Button(description="Filtern und Visualisieren", button_style='success',
                                   style={'button_color': accent_color})

    def on_filter_click(b):
        seriennummer = seriennummer_input.value
        start_date = date_picker_start.value
        end_date = date_picker_end.value
        filtered_df = filter_data(df, seriennummer=seriennummer, start_date=start_date, end_date=end_date)
        if filtered_df.empty:
            with container:
                container.clear_output()
                display("Keine Daten im angegebenen Zeitraum verfügbar.")
        else:
            plot_data(filtered_df)

    filter_button.on_click(on_filter_click)

    # Widgets anzeigen
    display(seriennummer_input, date_picker_start, date_picker_end, filter_button, container)


# Produktmanagement - Ländervergleich mit Nutzeranalyse auf Europakarten


# Produktmanagement - Produktausnutzung
def produktausnutzung_dashboard(df):
    container = widgets.Output()

    def plot_data(filtered_df):
        with container:
            container.clear_output()

            # Überprüfen, ob 'Start Time' vorhanden ist
            if 'Start Time' not in filtered_df.columns:
                display("Die Spalte 'Start Time' ist nicht im DataFrame vorhanden.")
                return

            # Konvertierung der 'Start Time' Spalte in das Datetime-Format
            filtered_df['Start Time'] = pd.to_datetime(filtered_df['Start Time'], errors='coerce', dayfirst=True)

            # Extrahierung des Wochentags und der Stunde
            filtered_df['Weekday'] = filtered_df['Start Time'].dt.day_name()

            # Extrahierung und Gruppierung der Stunde in 2-Stunden-Intervalle
            filtered_df['HourGroup'] = (filtered_df['Start Time'].dt.hour // 2) * 2
            filtered_df['HourGroup'] = filtered_df['HourGroup'].astype(str) + '-' + (filtered_df['HourGroup'] + 1).astype(str) + 'h'

            # Reihenfolge der Wochentage festlegen
            weekday_order = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
            filtered_df['Weekday'] = pd.Categorical(filtered_df['Weekday'], categories=weekday_order, ordered=True)

            # Gruppierung der Daten nach Wochentag und 2-Stunden-Intervallen und Zählen der Sessions
            df_grouped = filtered_df.groupby(['Weekday', 'HourGroup']).size().reset_index(name='Sessions')

            # Entfernen von Zeilen mit Nullwerten
            df_grouped = df_grouped[df_grouped['Sessions'] > 0]

            # Erstellen des Treemap-Diagramms
            fig = px.treemap(df_grouped, path=['Weekday', 'HourGroup'], values='Sessions', color='Sessions',
                             color_continuous_scale='RdYlBu_r', title='Treemap der Ladevorgänge (2-Stunden-Intervalle)')

            # CI-Anpassungen
            fig.update_layout(
                plot_bgcolor=secondary_color,
                paper_bgcolor=secondary_color,
                font_color=primary_color,
                title_font_color=primary_color,
                title_x=0.5,
                font=dict(size=14),
            )

            # Anzeige des Treemap-Diagramms
            fig.show()

    # Eingabefeld für die Seriennummer
    seriennummer_input = widgets.Text(
        value='',
        placeholder='Geben Sie die Seriennummer ein (optional)',
        description='Seriennummer:',
        style={'description_width': 'initial'},
        layout=widgets.Layout(width='300px')
    )

    # Datepicker für die Datumsfilterung
    date_picker_start = widgets.DatePicker(description='Startdatum:')
    date_picker_end = widgets.DatePicker(description='Enddatum:')

    # Button zum Filtern der Daten
    filter_button = widgets.Button(description="Filtern und Visualisieren", button_style='success',
                                   style={'button_color': accent_color})

    def on_filter_click(b):
        seriennummer = seriennummer_input.value
        start_date = date_picker_start.value
        end_date = date_picker_end.value
        filtered_df = filter_data(df, seriennummer=seriennummer, start_date=start_date, end_date=end_date)
        if filtered_df.empty:
            with container:
                container.clear_output()
                display("Keine Daten im angegebenen Zeitraum verfügbar.")
        else:
            plot_data(filtered_df)

    filter_button.on_click(on_filter_click)

    # Widgets anzeigen
    display(seriennummer_input, date_picker_start, date_picker_end, filter_button, container)

# Produktmanagement - Ländervergleich mit 4 interaktiven Europakarten
def laendervergleich_dashboard(df):
    container = widgets.Output()

    def plot_europe_map(filtered_df, title):
        fig = px.choropleth(
            filtered_df,
            locations='Country', 
            locationmode='country names', 
            color='Sessions',
            title=title,
            scope='europe',
            color_continuous_scale=px.colors.sequential.Blues
        )
        fig.update_layout(
            plot_bgcolor=secondary_color,
            paper_bgcolor=secondary_color,
            font_color=primary_color,
            title_font_color=primary_color,
            title_x=0.5,
            font=dict(size=14),
        )
        fig.show()

    # Beispieldaten für Österreich, Deutschland und Rumänien
    seriennummer_osterreich = '19BZ00401'
    seriennummer_deutschland = '23BZ2234B'
    seriennummer_rumaenien = '20BZ00001'

    # Daten nach Ländern filtern
    filtered_df_osterreich = filter_data(df, seriennummer=seriennummer_osterreich)
    filtered_df_deutschland = filter_data(df, seriennummer=seriennummer_deutschland)
    filtered_df_rumaenien = filter_data(df, seriennummer=seriennummer_rumaenien)

    # Erstellen von fiktiven Daten zu den Sessions für die Karten
    country_data = pd.DataFrame({
        'Country': ['Austria', 'Germany', 'Romania'],
        'Sessions': [
            filtered_df_osterreich.shape[0],
            filtered_df_deutschland.shape[0],
            filtered_df_rumaenien.shape[0]
        ]
    })

    # Titel für die Karten
    title_osterreich = f'Charger Sessions in Österreich (Seriennummer {seriennummer_osterreich})'
    title_deutschland = f'Charger Sessions in Deutschland (Seriennummer {seriennummer_deutschland})'
    title_rumaenien = f'Charger Sessions in Rumänien (Seriennummer {seriennummer_rumaenien})'

    # Interaktive Europakarten anzeigen
    with container:
        container.clear_output()
        plot_europe_map(country_data, title='Sessions in den Ländern')
        plot_europe_map(pd.DataFrame({'Country': ['Austria'], 'Sessions': [country_data['Sessions'][0]]}), title_osterreich)
        plot_europe_map(pd.DataFrame({'Country': ['Germany'], 'Sessions': [country_data['Sessions'][1]]}), title_deutschland)
        plot_europe_map(pd.DataFrame({'Country': ['Romania'], 'Sessions': [country_data['Sessions'][2]]}), title_rumaenien)

    display(container)

# Funktion zum Anzeigen des Hauptmenüs
def show_main_dashboard(mode='Analysis'):
    df = df_global

    # Dashboard-Titel basierend auf dem Modus festlegen
    if mode == 'Analysis':
        dashboard_title = 'MOON Analysis Dashboard'
    elif mode == 'Live':
        dashboard_title = 'MOON Live Dashboard'
    elif mode == 'Predictive':
        dashboard_title = 'MOON Predictive Dashboard'
    else:
        dashboard_title = 'MOON Dashboard'

    # Titel anzeigen
    display(widgets.HTML(f"<h2>{dashboard_title}</h2>"))

    # Aktualisierungsbutton
    refresh_button = widgets.Button(description="Aktualisieren", button_style='info',
                                    style={'button_color': accent_color})
    refresh_button.on_click(lambda b: refresh_dashboard())
    display(refresh_button)

    # Dropdown für Abteilungen
    dropdown_abteilung = widgets.Dropdown(
        options=abteilungen,
        value='Vertrieb',
        description='Abteilung:',
        style={'description_width': 'initial'},
    )
    display(dropdown_abteilung)

    # Dropdown für Subkarten
    dropdown_subkarte = widgets.Dropdown(
        options=subkarten['Vertrieb'],
        value=subkarten['Vertrieb'][0],
        description='Subkarte:',
        style={'description_width': 'initial'},
    )
    display(dropdown_subkarte)

    # Container für Dashboard-Anzeige
    container = widgets.Output()
    display(container)

    # Funktion zur Aktualisierung des Subkarten-Dropdowns
    def on_abteilung_change(change):
        dropdown_subkarte.options = subkarten[change['new']]
        dropdown_subkarte.value = subkarten[change['new']][0]
        on_subkarte_change(None)

    dropdown_abteilung.observe(on_abteilung_change, names='value')

    # **Hier die Funktion korrekt einrücken**
    def on_subkarte_change(change):
        with container:
            container.clear_output()
            abteilung = dropdown_abteilung.value
            subkarte = dropdown_subkarte.value
            try:
                if abteilung == 'Vertrieb':
                    if subkarte == 'Chargerausnutzung':
                        vertrieb_chargerausnutzung(df)
                    elif subkarte == 'Chargeranalyse':
                        vertrieb_chargeranalyse(df)
                    elif subkarte == 'Standortanalyse':
                        vertrieb_standortanalyse(df)
                    elif subkarte == 'Fahrzeuganalyse':
                        fahrzeuganalyse_dashboard(df)
                    elif subkarte == 'Amortisationsrechnung':
                        amortisationsrechnung_dashboard()
                    else:
                        display(f"Subkarte '{subkarte}' wird bald hinzugefügt.")
                elif abteilung == 'Produktmanagement':
                    if subkarte == 'Produktausnutzung':
                        produktausnutzung_dashboard(df)
                    elif subkarte == 'Nutzeranalyse':
                        nutzeranalyse_dashboard(df)
                    elif subkarte == 'Ländervergleich':
                        laendervergleich_dashboard(df)
                    elif subkarte == 'Performance':
                        # Funktion für Performance-Dashboard aufrufen
                        performance_dashboard(df)
                    else:
                        display(f"Subkarte '{subkarte}' wird bald hinzugefügt.")
                elif abteilung == 'Marketing':
                    if subkarte == 'Kundenanalyse':
                        kundenanalyse_dashboard(df)
                    elif subkarte == 'Länderanalyse':
                        marketing_dashboard()
                    else:
                        display(f"Subkarte '{subkarte}' wird bald hinzugefügt.")
                elif abteilung == 'After Sales':
                    if subkarte == 'Fehlerquote':
                        fehlerquote_dashboard(df)
                    elif subkarte == 'Fehlerarten':
                        fehlerarten_dashboard(df)
                    elif subkarte == 'Charger Fehleranalyse':
                        # Falls Sie eine Funktion dafür haben
                        charger_fehleranalyse_dashboard(df)
                    else:
                        display(f"Subkarte '{subkarte}' wird bald hinzugefügt.")
                elif abteilung == 'Datenpflege':
                    if subkarte == 'Datenpflege':
                        datenpflege_dashboard()
                    else:
                        display(f"Subkarte '{subkarte}' wird bald hinzugefügt.")
                elif abteilung == 'Reporting':
                    if subkarte == 'Reporting':
                        reporting_dashboard()
                    else:
                        display(f"Subkarte '{subkarte}' wird bald hinzugefügt.")
                else:
                    display(f"Abteilung '{abteilung}' und Subkarte '{subkarte}' werden bald hinzugefügt.")
            except Exception as e:
                display(f"Ein Fehler ist aufgetreten: {e}")

    dropdown_subkarte.observe(on_subkarte_change, names='value')

    # Initiales Dashboard anzeigen
    on_subkarte_change(None)


# Funktion zum Anzeigen des Hauptmenüs
def show_main_menu():
    container = widgets.Output()

    # Buttons erstellen
    button_analysis = widgets.Button(description="MOON Analysis", button_style='primary',
                                     style={'button_color': accent_color}, layout=widgets.Layout(width='200px'))
    button_live = widgets.Button(description="MOON Live", button_style='primary',
                                 style={'button_color': accent_color}, layout=widgets.Layout(width='200px'))
    button_predictive = widgets.Button(description="MOON Predictive", button_style='primary',
                                       style={'button_color': accent_color}, layout=widgets.Layout(width='200px'))
    button_full_control = widgets.Button(description="MOON Full control", button_style='primary',
                                         style={'button_color': accent_color}, layout=widgets.Layout(width='200px'))
    button_documents = widgets.Button(description="MOON Documents", button_style='primary',
                                      style={'button_color': accent_color}, layout=widgets.Layout(width='200px'))

    # Buttons anordnen
    buttons = widgets.HBox([button_analysis, button_live, button_predictive, button_full_control, button_documents],
                           layout=widgets.Layout(justify_content='space-between'))

    display(buttons, container)

    # Callback-Funktionen definieren
    def on_analysis_click(b):
        with container:
            container.clear_output()
            show_main_dashboard(mode='Analysis')

    def on_live_click(b):
        with container:
            container.clear_output()
            show_main_dashboard(mode='Live')

    def on_predictive_click(b):
        with container:
            container.clear_output()
            show_main_dashboard(mode='Predictive')

    def on_full_control_click(b):
        with container:
            container.clear_output()
            show_full_control_dashboard()

    def on_documents_click(b):
        with container:
            container.clear_output()
            show_documents_dashboard()

    # Callback-Funktionen mit den Buttons verknüpfen
    button_analysis.on_click(on_analysis_click)
    button_live.on_click(on_live_click)
    button_predictive.on_click(on_predictive_click)
    button_full_control.on_click(on_full_control_click)
    button_documents.on_click(on_documents_click)


# Funktion für das Anmeldebildschirm
def start_menu():
    container = widgets.Output()

     # Logo anzeigen
    display(Image(filename='moon_power_gmbh_logo.jpg', width=200))

    # Eingabefelder
    username_input = widgets.Text(
        value='',
        placeholder='Benutzername',
        description='Benutzername:',
        style={'description_width': 'initial'},
    )

    password_input = widgets.Password(
        value='',
        placeholder='Passwort',
        description='Passwort:',
        style={'description_width': 'initial'},
    )

    # Login-Button
    login_button = widgets.Button(description="Login", button_style='success',
                                  style={'button_color': accent_color})

    # Callback-Funktion für den Login-Button
    def login(b):
        with container:
            container.clear_output()
            username = username_input.value
            password = password_input.value

            if username == 'MOON POWER' and password == 'Moon':
                container.clear_output()
                show_main_menu()
            else:
                container.clear_output()
                display("Falscher Benutzername oder Passwort. Versuchen Sie es erneut.")

    login_button.on_click(login)

    # Widgets anzeigen
    display(username_input, password_input, login_button, container)

# Anmeldebildschirm starten
start_menu()

